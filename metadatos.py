import os
import sys
import docx
import openpyxl
import fitz

def extract_docx_metadata(file_path):
    doc = docx.Document(file_path)
    metadata = {
        "Title": doc.core_properties.title,
        "Author": doc.core_properties.author,
        "Subject": doc.core_properties.subject,
        "Keywords": doc.core_properties.keywords,
        "Comments": doc.core_properties.comments,
        "Category": doc.core_properties.category,
        "Content Status": doc.core_properties.content_status,
        "Language": doc.core_properties.language,
        "Last Modified By": doc.core_properties.last_modified_by,
        "Created": doc.core_properties.created,
        "Modified": doc.core_properties.modified,
        "Revision": doc.core_properties.revision,
        "Version": doc.core_properties.version,
        "Page Count": len(doc.element.body.xpath('.//w:lastRenderedPageBreak')),
        "Paragraph Count": len(doc.paragraphs),
        "Table Count": len(doc.tables),
        "Image Count": len(doc.inline_shapes),
        # Calculate Word Count and Character Count manually
        "Word Count": sum(len(p.text.split()) for p in doc.paragraphs),
        "Character Count": sum(len(p.text) for p in doc.paragraphs),
    }
    return metadata


def extract_xlsx_metadata(file_path):
    wb = openpyxl.load_workbook(file_path)
    metadata = {
        "Title": wb.properties.title,
        "Author": wb.properties.creator,
        "Subject": wb.properties.subject,
        "Keywords": wb.properties.keywords,
        "Comments": wb.properties.description,
        "Category": wb.properties.category,
        "Content Status": wb.properties.contentStatus,
        "Language": wb.properties.language,
        "Last Modified By": wb.properties.lastModifiedBy,
        "Created": wb.properties.created,
        "Modified": wb.properties.modified,
        "Revision": wb.properties.revision,
        "Version": wb.properties.version,
        "Sheet Names": wb.sheetnames,
        "Sheet Count": len(wb.sheetnames),
        "Table Count": sum(1 for ws in wb.worksheets for table in ws.tables),
        # Add more metadata as needed
    }
    return metadata

def extract_pdf_metadata(file_path):
    metadata = {}
    try:
        doc = fitz.open(file_path)
        metadata["Title"] = doc.metadata.get("title", "")
        metadata["Author"] = doc.metadata.get("author", "")
        metadata["Subject"] = doc.metadata.get("subject", "")
        metadata["Keywords"] = doc.metadata.get("keywords", "")
        metadata["Producer"] = doc.metadata.get("producer", "")
        metadata["Created"] = doc.metadata.get("creationDate", "")
        metadata["Modified"] = doc.metadata.get("modDate", "")
        metadata["Page Count"] = doc.page_count
        # Add more metadata as needed
    except Exception as e:
        print(f"Error al extraer metadatos del archivo PDF: {e}")
    return metadata


def extract_metadata(file_path):
    _, file_extension = os.path.splitext(file_path)
    file_extension = file_extension.lower()
    if file_extension == '.docx':
        return extract_docx_metadata(file_path)
    elif file_extension == '.xlsx':
        return extract_xlsx_metadata(file_path)
    elif file_extension == '.pdf':
        return extract_pdf_metadata(file_path)
    else:
        return None

def main(directory):
    for root, _, files in os.walk(directory):
        print("Processing directory:", root)
        for file in files:
            file_path = os.path.join(root, file)
            print("Processing file:", file_path)
            metadata = extract_metadata(file_path)
            if metadata:
                print("File:", file_path)
                for key, value in metadata.items():
                    print(f"{key}: {value}")
                print("\n")

if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Usage: python script.py <directory>")
        sys.exit(1)

    directory = sys.argv[1]
    main(directory)
